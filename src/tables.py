"""Создание сводных таблиц параметров ступенчатой решетки."""
import locale
from datetime import date
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.section import WD_ORIENT
from Dry.allgui import fstr, to_kw, to_mm
from Dry.allcalc import Distance
from stepscreen import (
    StepScreen, SCREEN_WIDTH_SERIES, SCREEN_HEIGHT_SERIES, InputData, START_DISCHARGE_FULL_HEIGHT,
    DIFF_TEETH_SERIES, TEETH_STEP_Y, THICKNESS_STEEL, STEEL_GAPS)
locale.setlocale(locale.LC_NUMERIC, '')


DISCHARGE_FULL_HEIGHT = {hs: START_DISCHARGE_FULL_HEIGHT + DIFF_TEETH_SERIES[hs] * TEETH_STEP_Y
                         for hs in SCREEN_HEIGHT_SERIES}


def _create_tables() -> None:
    document = Document()
    section = document.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(42)
    section.page_height = Cm(29.7)

    styles = document.styles
    styles['Normal'].font.name = 'Calibri'
    styles['Normal'].font.size = Pt(11)

    max_gap = STEEL_GAPS[-1]
    min_gap = STEEL_GAPS[0]
    moving_steel_s0, fixed_steel_s0 = THICKNESS_STEEL[0]
    moving_steel_s1, fixed_steel_s1 = THICKNESS_STEEL[-1]
    discharge_height = 1.0

    document.add_paragraph().add_run(
        'Диапазоны значений ниже приведены от самого легкого исполнения '
        '({}/{}, прозор {}) до самого тяжелого ({}/{}, прозор {}).\n'
        'Высота сброса принята {} м над каналом, материал пластин — сталь и полипропилен.'.format(
            fstr(to_mm(fixed_steel_s1)), fstr(to_mm(moving_steel_s1)), fstr(to_mm(max_gap)),
            fstr(to_mm(fixed_steel_s0)), fstr(to_mm(moving_steel_s0)), fstr(to_mm(min_gap)),
            fstr(discharge_height)))

    paragraph = document.add_paragraph()
    paragraph.add_run('Масса решетки (кг)')
    paragraph.style = 'Heading 2'
    table1 = document.add_table(
        rows=len(SCREEN_WIDTH_SERIES) + 1, cols=len(SCREEN_HEIGHT_SERIES) + 1)

    paragraph = document.add_paragraph()
    paragraph.add_run('Мощность привода (кВт)')
    paragraph.style = 'Heading 2'
    table2 = document.add_table(
        rows=len(SCREEN_WIDTH_SERIES) + 1, cols=len(SCREEN_HEIGHT_SERIES) + 1)

    for table in (table1, table2):
        for i, width_serie in enumerate(SCREEN_WIDTH_SERIES, start=1):
            table.cell(i, 0).text = '{}xx'.format(fstr(width_serie, '%02d'))
        for i, height_serie in enumerate(SCREEN_HEIGHT_SERIES, start=1):
            table.cell(0, i).text = 'xx{}'.format(fstr(height_serie, '%02d'))
        table.style = 'Light Grid Accent 1'
    for row, width_serie in enumerate(SCREEN_WIDTH_SERIES, start=1):
        for col, height_serie in enumerate(SCREEN_HEIGHT_SERIES, start=1):
            screen = (
                StepScreen(InputData(
                    screen_ws=width_serie, screen_hs=height_serie, main_steel_gap=max_gap,
                    fixed_steel_s=fixed_steel_s1, moving_steel_s=moving_steel_s1,
                    channel_height=Distance(
                        DISCHARGE_FULL_HEIGHT[height_serie] - discharge_height),
                    have_plastic_part=True)),
                StepScreen(InputData(
                    screen_ws=width_serie, screen_hs=height_serie, main_steel_gap=min_gap,
                    fixed_steel_s=fixed_steel_s0, moving_steel_s=moving_steel_s0,
                    channel_height=Distance(
                        DISCHARGE_FULL_HEIGHT[height_serie] - discharge_height),
                    have_plastic_part=True)),)
            table1.cell(row, col).text = '{}—{}'.format(
                fstr(screen[0].full_mass, '%.0f'),
                fstr(screen[1].full_mass, '%.0f'))
            table2.cell(row, col).text = '{}—{}'.format(
                fstr(to_kw(screen[0].drive_unit.power)) if screen[0].drive_unit is not None
                else 'none',
                fstr(to_kw(screen[1].drive_unit.power)) if screen[1].drive_unit is not None
                else 'none',
                )
    document.save('Мощность привода и масса РСК ({}).docx'.format(date.today()))


GAP_VAR = 'gap'
SS_VAR = 'ss'


def _power_and_mass(screen: StepScreen) -> str:
    desc = screen.description
    if screen.drive_unit is not None:
        power = fstr(to_kw(screen.drive_unit.power))
    else:
        power = '...'
    mass = screen.full_mass
    return(f'  Powers.Add CreateKey({GAP_VAR}, {SS_VAR}, "{desc}"), "{power}"\n'
           f'  Weights.Add CreateKey({GAP_VAR}, {SS_VAR}, "{desc}"), "{mass:.0f}"\n')


def _create_sheet() -> None:
    with open('CopyToCalc.bas', 'w') as file:
        for gap in STEEL_GAPS:
            for thickness_pair in THICKNESS_STEEL:
                moving_steel_s, fixed_steel_s = thickness_pair
                gap_f = '{:.0f}'.format(to_mm(gap))
                fix_f = '{:g}'.format(to_mm(fixed_steel_s))
                mov_f = '{:g}'.format(to_mm(moving_steel_s))
                file.write(f'Private Sub Init_Gap{gap_f}_{fix_f}{mov_f}()\n'
                           f'    Const {GAP_VAR} as String = "{gap_f}"\n'
                           f'    Const {SS_VAR} as String = "{fix_f}/{mov_f}"\n')
                for width_serie in SCREEN_WIDTH_SERIES:
                    for height_serie in SCREEN_HEIGHT_SERIES:
                        screen = StepScreen(InputData(
                            screen_ws=width_serie, screen_hs=height_serie, main_steel_gap=gap,
                            fixed_steel_s=fixed_steel_s, moving_steel_s=moving_steel_s,
                            channel_height=Distance(DISCHARGE_FULL_HEIGHT[height_serie] - 1.0),
                            have_plastic_part=True))
                        file.write(_power_and_mass(screen))
                file.write('End Sub\n')


if __name__ == '__main__':
    _create_tables()
    _create_sheet()
